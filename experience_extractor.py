# app.py
import streamlit as st
import os
import re
import json
import tempfile
import time
from datetime import datetime
from typing import Optional, Tuple, List

# ---------- Page config & Styling ----------
st.set_page_config(page_title="Experience Extractor", layout="wide")

st.markdown(
    """
    <style>
    /* Page / card styling */
    .card {
        background: linear-gradient(180deg, rgba(255,255,255,0.02), rgba(255,255,255,0.01));
        border-radius: 14px;
        padding: 24px;
        margin-bottom: 18px;
        box-shadow: 0 8px 28px rgba(2,8,23,0.6);
    }

    /* File title / resume name */
    .file-name {
        font-size: 30px;
        font-weight: 900;
        color: #7ed0ff;
        margin-bottom: 8px;
    }

    /* Experience label + value */
    .exp-label {
        font-size: 16px;
        color: #c7d0d6;
        margin-bottom: 6px;
    }
    .exp-value {
        font-size: 32px;
        font-weight: 900;
        color: #9ef3c9;
        margin-bottom: 10px;
    }

    .computed {
        font-size: 14px;
        color: #9aa5b1;
        margin-top: 14px;
    }

    /* Make page containers full width feel */
    .app-body {
        padding-left: 18px;
        padding-right: 18px;
    }

    /* Button styling */
    div.stButton > button {
        background-color:#2E8B57;
        color:white;
        font-size:15px;
        padding:10px 18px;
        border-radius:8px;
    }
    div.stButton > button:hover {
        background-color:#3CB371;
    }

    /* Responsive tweaks */
    @media (max-width: 768px) {
        .file-name { font-size: 22px; }
        .exp-value { font-size: 22px; }
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# ---------- Optional dependencies ----------
try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    import docx
except Exception:
    docx = None

# ---------- Optional Gemini client ----------
GEMINI_CLIENT = None
try:
    from google import genai as _genai
    GEMINI_CLIENT = _genai
except Exception:
    GEMINI_CLIENT = None

# ---------- Key retrieval ----------
def get_gemini_key() -> str:
    try:
        if "GEMINI_API_KEY" in st.secrets:
            return st.secrets["GEMINI_API_KEY"]
    except Exception:
        pass
    return os.getenv("GEMINI_API_KEY", "").strip()

# ---------- Text extraction helpers ----------
def extract_text_from_pdf(path: str) -> str:
    if pdfplumber is None:
        raise RuntimeError("Install pdfplumber: pip install pdfplumber")
    parts = []
    with pdfplumber.open(path) as pdf:
        for p in pdf.pages:
            parts.append(p.extract_text() or "")
    return re.sub(r'\s+', ' ', " ".join(parts)).strip()

def extract_text_from_docx(path: str) -> str:
    if docx is None:
        raise RuntimeError("Install python-docx: pip install python-docx")
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    parts.append(cell.text.strip())
    return re.sub(r'\s+', ' ', " ".join(parts)).strip()

def extract_text_from_txt(path: str) -> str:
    with open(path, "rb") as f:
        raw = f.read()
    return re.sub(r'\s+', ' ', raw.decode('utf-8', errors='ignore')).strip()

def extract_text(path: str) -> str:
    p = path.lower()
    if p.endswith(".pdf"):
        return extract_text_from_pdf(path)
    if p.endswith(".docx"):
        return extract_text_from_docx(path)
    if p.endswith(".txt"):
        return extract_text_from_txt(path)
    raise ValueError("Unsupported file type. Supported: .pdf, .docx, .txt")

# ---------- Heuristic fallback for years ----------
def fallback_years(text: str) -> float:
    vals = []
    if not text:
        return 0.0
    for m in re.finditer(r'(\d+(?:\.\d+)?)\s*(?:\+)?\s*(?:years?|yrs?)', text.lower()):
        try:
            vals.append(float(m.group(1)))
        except:
            pass
    return round(max(vals), 1) if vals else 0.0

# ---------- Gemini LLM call for years ----------
def build_years_prompt(text: str, max_chars=12000) -> str:
    if len(text) > max_chars:
        text = text[:max_chars]
    today = datetime.now().strftime("%Y-%m-%d")
    prompt = f"""
Return ONLY JSON: {{ "total_years": <float> }}.

Compute TOTAL professional work experience:
- Merge overlapping roles.
- Convert months into decimal years (1 decimal).
- Treat "present/current" as {today}.
- If unsure → return 0.0
- DO NOT output anything except JSON.

Resume:
\"\"\"{text}\"\"\"
""".strip()
    return prompt

def ask_gemini_for_years(text: str, api_key: str, max_retries: int = 2) -> float:
    if not text:
        return 0.0
    prompt = build_years_prompt(text)
    # fallback if client or key missing
    if not api_key or GEMINI_CLIENT is None:
        return fallback_years(text)

    # Build client (try common variants)
    client = None
    try:
        client = GEMINI_CLIENT.Client(api_key=api_key)
    except Exception:
        try:
            GEMINI_CLIENT.configure(api_key=api_key)
            client = GEMINI_CLIENT
        except Exception:
            client = None

    if client is None:
        return fallback_years(text)

    for attempt in range(max_retries + 1):
        try:
            resp = client.models.generate_content(model="gemini-2.5-flash", contents=prompt)
            raw = resp.text if hasattr(resp, "text") else str(resp)
            m = re.search(r'\{.*\}', raw, flags=re.DOTALL)
            if m:
                try:
                    obj = json.loads(m.group(0))
                    val = float(obj.get("total_years", 0.0))
                    return round(val, 1)
                except Exception:
                    pass
            # try to parse first number if JSON not returned
            m2 = re.search(r'(\d+(?:\.\d+)?)', raw)
            if m2:
                return round(float(m2.group(1)), 1)
            return fallback_years(text)
        except Exception:
            if attempt < max_retries:
                time.sleep(0.6 * (attempt + 1))
                continue
            return fallback_years(text)

# ---------- Convert decimal years to human-readable ----------
def convert_decimal_to_human(decimal_years: float, method: str = "round") -> str:
    years = int(decimal_years)
    if method == "floor":
        months = int((decimal_years - years) * 12)
    else:
        months = int(round((decimal_years - years) * 12))

    # adjust overflow
    if months >= 12:
        years += 1
        months -= 12

    if years == 0 and months == 0:
        return "0 years"
    if months == 0:
        return f"{years} years"
    if years == 0:
        return f"{months} months"
    return f"{years} years {months} months"

# ---------- UI: Title & instructions ----------
st.title("🎯 Experience Extractor — Resume Batch")
st.write(
    "Upload resumes (.pdf / .docx / .txt). The app will attempt to use Gemini if "
    "`GEMINI_API_KEY` is set in `st.secrets` or environment. Otherwise a heuristic fallback is used."
)

api_key = get_gemini_key()
if api_key:
    st.success("Gemini API key loaded — model extraction enabled.")
else:
    st.info("No Gemini API key found — using local heuristic fallback. To enable model extraction, add `GEMINI_API_KEY` in Streamlit Secrets or environment variables.")

# ---------- Upload control ----------
uploaded = st.file_uploader("Upload resumes", type=["pdf", "docx", "txt"], accept_multiple_files=True)

if uploaded:
    # Button to start processing
    if st.button("🔎 Extract Experience from Uploads"):
        results = {}
        with st.spinner("Processing resumes — extracting text and computing experience..."):
            for f in uploaded:
                suffix = os.path.splitext(f.name)[1]
                with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                    tmp.write(f.read())
                    tmp_path = tmp.name
                try:
                    text = extract_text(tmp_path)
                    decimal = ask_gemini_for_years(text, api_key)
                    human = convert_decimal_to_human(decimal, method="round")
                    results[f.name] = {"decimal": decimal, "human": human}
                except Exception as e:
                    results[f.name] = {"error": str(e)}
                finally:
                    try:
                        os.remove(tmp_path)
                    except Exception:
                        pass

        # ---------- Display FULL-WIDTH cards (only human-readable shown) ----------
        st.markdown("<hr/>", unsafe_allow_html=True)
        st.markdown("## 🟩 Results (Experience Extracted)")

        for fname, out in results.items():

            if "error" in out:
                st.markdown(
                    f"""
                    <div class="card" style="width:100%;">
                        <div class="file-name">{fname}</div>
                        <div style="color:#ffaaaa; font-size:16px;">Error: {out["error"]}</div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
                continue

            human = out.get("human", "0 years")

            html = f"""
            <div class="card" style="width:100%;">
                <div class="file-name">{fname}</div>

                <div>
                    <div class="exp-label">Experience</div>
                    <div class="exp-value">{human}</div>
                </div>

                <div class="computed">Computed at {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</div>
            </div>
            """

            st.markdown(html, unsafe_allow_html=True)


    # end of if button
else:
    st.info("Upload one or more resume files to get started.")
